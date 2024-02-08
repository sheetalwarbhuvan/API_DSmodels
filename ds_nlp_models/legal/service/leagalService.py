 
from re import A
from docx import Document
from datetime import datetime

from langchain_community import vectorstores
# from langchain_community.schema import prompt
import os

from PyPDF2 import PdfReader
from langchain.memory import ConversationBufferMemory
from langchain.chains.conversation.memory import ConversationBufferMemory
from langchain.chains import ConversationalRetrievalChain

from langchain.text_splitter import CharacterTextSplitter
from langchain_community.embeddings import HuggingFaceInstructEmbeddings
from langchain_community.vectorstores import FAISS
from langchain_community.llms import HuggingFaceHub
from langchain_community.document_loaders import PyPDFLoader
from transformers import AutoTokenizer, AutoModelForSeq2SeqLM

from langchain_community.llms import HuggingFaceHub
from transformers import pipeline
import numpy as np


import pandas as pd
from langchain_community.document_loaders import PyPDFLoader
import re
from langchain.schema.output import LLMResult
from transformers import T5Tokenizer, T5ForConditionalGeneration,AutoTokenizer

from_tf=True
# Function to generate "Statement of Work" text based on user inputs
def generate_statement_of_work(company_name, vendor_name, effective_date, agreement_date):
    statement_of_work = (
        f"This Statement of Work (SOW) is made on {effective_date}, by and between\n"
        f"{company_name} (hereinafter referred to as the 'Company') and {vendor_name} "
        f"(hereinafter referred to as the 'Vendor'), individually a 'Party' and collectively as 'Parties'. "
        f"The SOW is effective from {effective_date} and shall be governed by the terms and conditions "
        f"of the Master Service Agreement (hereinafter referred to as the 'Agreement') between {company_name} "
        f"and {vendor_name} dated {agreement_date}. Terms used in this SOW that are not defined herein shall "
        f"have the meanings given to them in the Agreement."
    )
    return statement_of_work




def documentCreation(data):
         # Define prompts for each topic
         company_name=data['company_name']
         effective_date=data['effective_date']
         vendor_name=data['vendor_name']
         agreement_date=data['agreement_date']
         businessChallenge=data['businessChallenge']
         project_initiation_fee=data['project_initiation_fee']
         monthly_service_fee=data['monthly_service_fee']
         company_representative_name=data['company_representative_name']
         vendor_representative_name=data['vendor_representative_name']
         company_representative_title=data['company_representative_title']
         vendor_representative_title=data['vendor_representative_title']
        
         prompts = {
            "Business Challenges": "Please provide an overview of the business challenges faced by the Company.",
            "Solution Overview": "Please provide an overview of the proposed solution to address the Business Challenges mentioned.",
            "Functional Scope": "Explain the functional scope of the project in detail, including any specific features and capabilities.",
            "Project Deliverables": "List and describe the key deliverables that will be provided as part of this project.",
            "Estimated Timelines": "Provide high-level estimated Timelines for completing this project.",
            "Risks": "Identify potential risks associated with the project and describe your mitigation strategy for each.",
            "Assumptions": "List any assumptions made for planning purposes and specify any factors considered true or certain.",
            "Onsite Requirements": "Outline the requirements and resources that the Customer will provide to the project team while onsite.",
            "Roles and Responsibilities": "Define the roles and responsibilities of the project team members and stakeholders.",
            "Payments": f"Billing rates for the team are detailed in Appendix B to this SOW for {company_name}, as mutually agreed upon by "
                        f"the Parties. Rates provided in Appendix B are indicative and are subject to change as mutually agreed by "
                        f"the Parties. Please provide any additional payment instructions or details:",
            "Terms and Conditions": "Provide terms and conditions."
        }
         checkpoint = "MBZUAI/LaMini-Flan-T5-783M"
         # Initialize the tokenizer and base model for text generation
         tokenizer = AutoTokenizer.from_pretrained(checkpoint)
         model = AutoModelForSeq2SeqLM.from_pretrained(checkpoint)  
        # Generate "Statement of Work" text
         sow_text = generate_statement_of_work(company_name, vendor_name,effective_date, agreement_date)
         # Initialize a dictionary to store the generated content for each topic
         sow_report = {}
         # Generate content for each topic based on the prompts
         for topic, prompt in prompts.items():
             # Create a combined prompt including the user's input for "Business Challenges"
            full_prompt = f"{prompt}\n\n{data['businessChallenge']}\n\n{topic}:"

            # Tokenize the prompt and generate content
            input_ids = tokenizer(full_prompt, return_tensors="pt").input_ids
            output = model.generate(input_ids, max_length=500, num_return_sequences=1, temperature=0.7, do_sample=True)
        #     output = model.generate(input_ids, max_length=512, num_return_sequences=1, temperature=0.7, do_sample=True, max_new_tokens=1000, eos_token_id=-1)

            # Decode and store the generated content
            generated_text = tokenizer.decode(output[0], skip_special_tokens=True)
            sow_report[topic] = generated_text
         # Create a new Word document
         doc = Document()
         doc.add_heading("AI Powered SOW Document",0)
         # Add content to the document, including tables for each topic
         for topic, text in sow_report.items():
            doc.add_heading(f'{topic}', level=1)

            # Add table for "Roles and Responsibilities"
            if topic == "Roles and Responsibilities":
                roles_table = doc.add_table(rows=1, cols=2)
                roles_table.style = 'Table Grid'
                roles_headers = roles_table.rows[0].cells
                roles_headers[0].text = 'Role'
                roles_headers[1].text = 'Responsibility'

                # Add sample data (replace with actual data)
                roles_data = [
                    ('Project Manager', 'Overall project coordination and management'),
                    ('Business Analyst', 'Gathering and documenting project requirements'),
                    ('UI/UX Designer', 'Designing user interfaces and user experiences'),
                    ('Development Team', 'Implementing core system functionalities'),
                    ('QA Team', 'Conducting unit testing for developed features')
                ]

                for row_data in roles_data:
                    row_cells = roles_table.add_row().cells
                    for i, cell_data in enumerate(row_data):
                        row_cells[i].text = cell_data
            else:
                # Add non-table content
                doc.add_paragraph(text)
        # Add the information about Project Sponsors and Vendor Fees
         doc.add_heading('Project Sponsors and Vendor Fees', level=1)
         doc.add_paragraph(f"Vendor Sponsor: {vendor_name}\nClient Sponsor: {company_name}\n\n"
                        f"Vendor Fees for the Services:\n"
                        f"Project Initiation Fee: ${project_initiation_fee} (payable upon project kick-off)\n"
                        f"Monthly Service Fee: ${monthly_service_fee} (payable at the beginning of each month)\n"
                        "Out-of-Pocket Expenses: Billed on a monthly basis, subject to pre-approval.")

        # Add the information about representatives
         doc.add_heading('FOR Company FOR Vendor', level=1)
         doc.add_paragraph(f"By: {company_representative_name}\t\t\t\tBy: {vendor_representative_name}\n"
                        f"Name: {company_representative_name}\t\t\t\tName: {vendor_representative_name}\n"
                        f"Title: {company_representative_title}\t\t\tTitle: {vendor_representative_title}")


        # Add Appendix A with table
         doc.add_heading('Appendix A: Project Deliverables and Timelines', level=1)
         appendix_a_table = doc.add_table(rows=1, cols=4)
         appendix_a_table.style = 'Table Grid'
         appendix_a_headers = appendix_a_table.rows[0].cells
         appendix_a_headers[0].text = 'Phase'
         appendix_a_headers[1].text = 'Deliverables'
         appendix_a_headers[2].text = 'Owner'
         appendix_a_headers[3].text = 'Timeline'
         
         # Add data to Appendix A table
         appendix_a_data = [
            ('1', 'Project Initiation Document (PID)', 'Project Manager', 'Within 2 weeks'),
            ('2', 'System Requirements Specification', 'Business Analyst', 'Within 4 weeks'),
            ('3', 'Prototype Design', 'UI/UX Designer', 'Within 6 weeks'),
            ('4', 'Development Phase', 'Development Team', 'Within 8 weeks'),
            ('5', 'User Acceptance Testing (UAT)', 'QA Team', 'Within 2 weeks')
        ]

         for row_data in appendix_a_data:
            row_cells = appendix_a_table.add_row().cells
            for i, cell_data in enumerate(row_data):
                row_cells[i].text = cell_data

        # Add Appendix B with table
         doc.add_heading('Appendix B: Rate Card', level=1)
         appendix_b_table = doc.add_table(rows=1, cols=4)
         appendix_b_table.style = 'Table Grid'
         appendix_b_headers = appendix_b_table.rows[0].cells
         appendix_b_headers[0].text = 'Verticals'
         appendix_b_headers[1].text = 'Role'
         appendix_b_headers[2].text = 'Location (Onsite/Remote)'
         appendix_b_headers[3].text = 'Headcount Rate (US$ per hour)'

        # Add data to Appendix B table
         appendix_b_data = [
            ('Development', 'Senior Developer', 'Onsite', '2 $80'),
            ('Testing', 'QA Engineer', 'Onsite', '1 $60'),
            ('Design', 'UI/UX Designer', 'Remote', '2 $70')
        ]
         for row_data in appendix_b_data:
            row_cells = appendix_b_table.add_row().cells
            for i, cell_data in enumerate(row_data):
                row_cells[i].text = cell_data
         return doc
         
         
def goldContract(file):
    ques = [
        'Is SOW effective date mentioned in the document?',
        'Does Risk section includes information about techincal failure and security issues?',
        'Does the client assume full responsibility for any potential risks or challenges that may arise during the implementation, as stated in the document?',
        'Does the SOW explicitly declare compliance with relevant laws and regulations?',
        'Does the SOW include provisions for renewal or extension under specific conditions?',
        'Does the document specify the payment terms and conditions for the project?',
        'Does the SOW provide details about the dispute resolution process in case of conflicts?'
    ]

        # 'Is SOW effective date mentioned in the document?':'Summarize and Give me the importance of why SOW date should be mentioned in a contract and also give example sentence to be included in SOW document',
        # 'Does Risk section includes information about techincal failure and security issues?': 'Summarize and Give a Sample Risk section which includes information about techincal failure and security issues',
        # 'Does the client assume full responsibility for any potential risks or challenges that may arise during the implementation, as stated in the document?':'Summarize and Give me the importance and a sample statement about how the client assume full responsibility for any potential risks or challenges that may arise during the implementation',
        # 'Does the SOW explicitly declare compliance with relevant laws and regulations?':'Summarize and Give me the importance of including the declaration of compliance with relevant laws and regulations in SOW Document and a sample sentence about how SOW explicitly declare compliance with relevant laws and regulations to be included in SOW document',
        # 'Does the SOW include provisions for renewal or extension under specific conditions?':'Summarize and Give me the importance of including provisions for renewal or extension under specific conditions and give me a sample statement about how SOW can include provisions for renewal or extension under specific conditions',
        # 'Does the document specify the payment terms and conditions for the project?':'Summarize and Give me the importance of including payment terms and codnitions and also give me a statement about what all to be included in the document incase of the payment terms and conditions for the project',
        # 'Does the SOW provide details about the dispute resolution process in case of conflicts?':'Summarize and Give a statement about dispute resolution process in case of conflicts and also give a sample statement to be included in SOW document'
    
    missing_answer_prompt = {
        'Is SOW effective date mentioned in the document?':'Give me the importance of why SOW date should be mentioned in a contract within 18 tokens',
        'Does Risk section includes information about techincal failure and security issues?': 'Give a Sample Risk section which includes information about techincal failure and security issues',
        'Does the client assume full responsibility for any potential risks or challenges that may arise during the implementation, as stated in the document?':'Summarize and Give me the importance of why the client should assume full responsibility for any potential risks or challenges that may arise during the implementation',
        'Does the SOW explicitly declare compliance with relevant laws and regulations?':'Give me the importance of including the declaration of compliance with relevant laws and regulations in SOW Document',
        'Does the SOW include provisions for renewal or extension under specific conditions?':'Summarize and Give me the importance of including provisions for renewal or extension under specific conditions',
        'Does the document specify the payment terms and conditions for the project?':'Summarize and Give me the importance of including payment terms and codnitions in SOW Document',
        'Does the SOW provide details about the dispute resolution process in case of conflicts?':'Summarize and Give a statement about dispute resolution process in case of conflicts in SOW document'
    }

        # 'When was the SOW made effective?',
        # 'Which expenses will the Vendor bill according to the terms mentioned in the SOW? Specifically, mention any out-of-pocket expenses.',
        # 'Project duration?',
        # 'Could you offer detailed insights into the specific goals and scope outlined for the CRM project as specified in the Statement of Work (SOW)?',
        # 'What specific objectives does XYZ Corporation aim to achieve through the CRM project, as per the SOW?',
        # 'What are the key deliverables outlined in the SOW for the CRM project?',
        # 'Please provide detailed information on the billing rates for the team, as mentioned in the Rate Card of the SOW',
        # 'Can you share the details from the Rate Card in Appendix B, particularly focusing on the rates for different roles and verticals?'

    response,suggestion=get_answers_for_questions(file,ques,missing_answer_prompt)
    results_table_data = []
    for question, response, suggestion in zip(ques, response, suggestion):
            results_table_data.append((question, response, suggestion))

    results_df = pd.DataFrame(results_table_data, columns=["Clause", "Is it in PDF or not?", "Suggestion"])
    print(results_df)
    return results_df
    

DEFAULT_SYSTEM_PROMPT = """
You are a helpful, respectful and honest assistant. 
If a question does not make any sense, or is not factually coherent, explain why instead of answering something not correct. 
If you don't know the answer to a question, please don't share false information.
""".strip()
   
def llm_get_suggestion():
    huggingfacehub_api_token = "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"
    llm = HuggingFaceHub(
        repo_id="google/flan-t5-xxl",
        # model_kwargs={"temperature": 0.5, "max_length": 512},
        model_kwargs={"temperature": 0.5, "max_length": 512, "max_new_tokens":1000, "eos_token_id":-1},
        huggingfacehub_api_token=huggingfacehub_api_token
    )
    return llm

llm_suggestion_response = None  # Initialize with a default value

def contains_negative(response):
    negative_patterns = [
        r'\b(no|not|does not contain|unanswerable|no answer)\b',
        # Add more patterns as needed
    ]
    
    for pattern in negative_patterns:
        if re.search(pattern, response, re.IGNORECASE):
            return True
    return False

def llm_get_suggestion():
    huggingfacehub_api_token = "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"
    llm = HuggingFaceHub(
        repo_id="google/flan-t5-xxl",
        # model_kwargs={"temperature": 0.5, "max_length": 512},
        model_kwargs={"temperature": 0.5, "max_length": 512, "max_new_tokens":1000, "eos_token_id":-1},
        huggingfacehub_api_token=huggingfacehub_api_token
    )
    return llm


def get_answers_for_questions(file,ques,missing_answer_prompt):
    responses = []
    suggestions = []
    vectorestore=process_pdfs(file)
    # Get responses for all questions
    for question in ques:
        system_prompt = DEFAULT_SYSTEM_PROMPT
        # Reset the conversation history for a new question
      

        # create conversation chain
        conversation = get_conversation_chain(vectorestore)
        # response =  st.session_state.conversation({'question': ques})
        response =  conversation.run(question)
        # st.session_state.chat_history = response['chat_history']

        if contains_negative(response):
            # Use the llm to generate a suggestion
            llm_instance = llm_get_suggestion()
            llm_prompt = missing_answer_prompt[question]

            try:
                llm_suggestion_response = llm_instance.generate([llm_prompt])
                # st.write("ouput 0:",llm_suggestion_response.generations[0][0].text)

            except Exception as e:
                print(f"Error during language model execution: {e}")

            if isinstance(llm_suggestion_response, LLMResult):
                # Access the content of all generated outputs
                suggestion = [generation.text if generation and generation.text else 'N/A' for generations_list in llm_suggestion_response.generations for generation in generations_list]

            else:
                suggestion = ['No Suggestion']

        else:
            suggestion = ['Information is Present']

        suggestions.append(suggestion[0])   
        responses.append(response)

    return responses, suggestions
def get_vectorstore(text_chunks):
    # st.write("TEXT CHUNKS",text_chunks)
    embeddings = HuggingFaceInstructEmbeddings(model_name="hkunlp/instructor-xl")
    vectorstore = FAISS.from_texts(texts=text_chunks, embedding=embeddings)
    # Persist the vectors locally on disk
    vectorstore.save_local("faiss_index_constitution")
    # Load from local storage
    persisted_vectorstore = FAISS.load_local("faiss_index_constitution", embeddings)
    return persisted_vectorstore
def process_pdfs(pdf_docs):
    raw_text = get_pdf_text(pdf_docs)
    text_chunks = get_text_chunks(raw_text)
    vectorstore = get_vectorstore(text_chunks)
    
    return vectorstore
def get_pdf_text(pdf_docs):
    text = ""
    for pdf in pdf_docs:
        pdf_reader = PdfReader(pdf)
        for page in pdf_reader.pages:
            text += page.extract_text()
    return text
    
def get_text_chunks(text):
    text_splitter = CharacterTextSplitter(
        separator="\n",
        chunk_size=1000,
        chunk_overlap=200,
        length_function=len
    )
    chunks = text_splitter.split_text(text)
    return chunks

def get_conversation_chain(persisted_vectorstore):
    huggingfacehub_api_token = "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"
    llm = HuggingFaceHub(
        # repo_id="MBZUAI/LaMini-Flan-T5-783M",
        repo_id="google/flan-t5-xxl",
        # model_kwargs={"temperature": 0.5, "max_length": 512},
        
        # model_kwargs={"temperature": 0.5,"max_length": 1024},
        # model_kwargs={"temperature": 1,"max_length": 1000000},
        model_kwargs={"temperature": 0.5, "max_length": 1000000, "max_new_tokens":1000000, "eos_token_id":-1},
        huggingfacehub_api_token=huggingfacehub_api_token
    )
    memory = ConversationBufferMemory(
        memory_key='chat_history', return_messages=True)
    conversation_chain = ConversationalRetrievalChain.from_llm(
        llm=llm,
        retriever=persisted_vectorstore.as_retriever(),
        memory=memory
    )
    return conversation_chain
